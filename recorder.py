#!/usr/bin/env python 
# -*- coding: utf-8 -*-
# @Time    : 2021/8/15 11:17
# @Author  : hanyou
# @Software: PyCharm
"""
用于实现打卡功能
"""
# 新事项出现时，其后5分钟（倒计时）弹出旧事项打卡选项，休息日统计分数


import json
from tkinter import *

from apscheduler.schedulers.blocking import BlockingScheduler
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tkinter
import tkinter.messagebox
from docx import Document

from static_parameters import s_param

with open('json_files/d_parameters_day.json', 'r') as f:
    d_h_param = json.load(f)

with open('json_files/d_parameters_sec.json', 'r') as f:
    d_param = json.load(f)

# 判读是否弹出窗口
def record_trigger():
    with open('json_files/d_parameters_sec.json', 'r') as f:
        d_param = json.load(f)
    if not (d_param['current_time_index'] == 0 and (
            d_h_param['day'] == s_param['genTable_next_month_start_day'] or d_h_param['day'] == s_param[
        'genTable_current_month_start_day'])):

        #安装时间表定时执行
        schedudler = BlockingScheduler()
        for i in range(s_param['things_len']):#只执行一次
            schedudler.add_job(window_pop, 'date', run_date=d_h_param['current_date_str_2']+' '+ s_param['schedule_time_list'] [i])
            # if i== s_param['things_len']-1:
            #     schedudler.shutdown(wait=False)#False立刻中断
        schedudler.start()

def window_pop():
    root = tkinter.Tk()
    var = tkinter.IntVar()
    # 按时间表弹出窗口
    lb = tkinter.Label(root, text='%s' % d_param['last_thing_str'], fg='black', font=("黑体", 50))
    lb.pack()
    # 值为1
    choose_yes = tkinter.Radiobutton(root, text="yes", variable=var, value=s_param['task_complete_value'],
                                     font=("cambria", 20))
    choose_yes.pack()
    # 值为2
    choose_no = tkinter.Radiobutton(root, text="no", variable=var, value=s_param['task_not_complete_value'],
                                    font=("cambria", 20))
    choose_no.pack()

    def write_result():
        num = var.get()
        # 如果值为1,写入
        # 获取当前单元格位置
        if num == s_param['task_complete_value']:
            f = Document(s_param['file_path'] + '%s' % d_h_param['current_headline_str'] + s_param['file_tpye'])
            table = f.tables[0]
            data_list = [int(table.cell(x + 1, 0).text[-2:]) for x in range(s_param['days_len'])]
            if d_param['last_time_index'] == s_param['things_len']:  # 最后一项
                write_row = data_list.index(d_h_param['day'])
            else:
                write_row = data_list.index(d_h_param['day']) + 1
            table.cell(write_row, d_param['last_time_index'] + 1).paragraphs[0].text = '✔'
            table.cell(write_row, d_param['last_time_index'] + 1).paragraphs[
                0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            f.save(s_param['file_path'] + '%s' % d_h_param['current_headline_str'] + s_param['file_tpye'])  # 需要保存！

        # 关闭窗口
        root.quit()

    # 按钮
    submit = Button(root, text="Submit", command=write_result)
    submit.pack()
    root.mainloop()


if __name__ == '__main__':
    pass
