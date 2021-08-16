#!/usr/bin/env python 
# -*- coding: utf-8 -*-
# @Time    : 2021/8/7 18:04
# @Author  : hanyou
# @Software: PyCharm
"""
用于两个时段生成表格
"""
import json

from docx import Document
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 重点

from static_parameters import s_param

with open('json_files/d_parameters_day.json', 'r') as f:
    d_h_param = json.load(f)


def creat_table():
    document = Document()
    # 转横向
    section = document.sections[0]  # 第一页设置横向
    # section = document.add_section(start_type=WD_SECTION_START.CONTINUOUS) # 添加横向页的连续节，多加一页
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE  # 设置横向
    section.page_width = new_width  # 设置宽高后才可
    section.page_height = new_height

    # 设置页边距
    section.top_margin = Cm(0.58)
    section.bottom_margin = Cm(1.18)
    section.left_margin = Cm(1.54)
    section.right_margin = Cm(1.54)

    # 设置字体颜色等
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 与上面合起来才有效果
    document.styles['Normal'].font.size = Pt(10.5)  # 五号
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    # 添加表头
    genTable_headline_str = d_h_param['genTable_headline_str']
    p = document.add_paragraph()
    head_line = p.add_run('%s' % genTable_headline_str)
    head_line.font.size = Pt(22)  # 二号
    head_line.font.name = u'黑体'  # 单独设置headline的字体
    head_line._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    p.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
    head_line.bold = True  # 字体加粗

    # 添加表格
    table = document.add_table(rows=s_param['table_row_len'], cols=s_param['table_col_len'],
                               style='Table Grid')  # 16行13列,正常样式
    # table.cell(0, 0).merge(table.cell(2, 2))  可以合并单元格

    # 设置列宽
    cols_width_list = s_param['col_width_list']
    for i in range(s_param['table_col_len']):
        table.cell(0, i).width = Cm(cols_width_list[i])

    # 设置行高
    for i in range(s_param['table_row_len'] - 2):  # 最后一行与第一行单独设置16-2
        table.rows[i + 1].height = Cm(0.9)
    table.rows[0].height = Cm(1)
    table.rows[s_param['table_row_len']-1].height = Cm(2.82)#row 10是第11行了要减1

    # 垂直居中对其
    # table.alignment = WD_TABLE_ALIGNMENT.CENTER#这个是单元格的对齐，不是内容的对其
    for i in range(s_param['table_row_len']):
        for j in range(s_param['table_col_len']):
            table.cell(i, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 这里全局设置居中后后面就不用设置了
            table.cell(i, j).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 添加事项
    table_col_list = s_param['table_col_list']
    for i in range(s_param['table_col_len']):
        row_head = table.rows[0].cells[i].paragraphs[0]
        row_head.text = table_col_list[i]
        # p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  #单独设置居中的时候可以这样做

    # 添加备注
    bak_list = s_param['bak_list']
    for i in range(s_param['table_col_len']):
        row_end = table.rows[-1].cells[i].paragraphs[0]
        row_end.text = bak_list[i]

    # 添加表格日期
    # 检测是15号生成本月，月底前第三天生成下月
    # 月底时间28 28-15=13,需要显示14天则14-13=1天划掉
    if d_h_param['genTable_switch'] == s_param['gentable_current_mode_num'] or \
            d_h_param['genTable_switch'] == s_param['gentable_repaire_later_mode_num']:
        remain_days_count = d_h_param['normal_genTable_later_day'] - d_h_param['normal_genTable_early_day']
        blank_days_count = s_param['days_len'] - remain_days_count  # 最少也会空一天
        for i in range(remain_days_count):  # 正
            col_0 = table.cell(i + 1, 0).paragraphs[0]
            col_0.text = str(d_h_param['month']) + '.' + '%s' % (s_param['genTable_current_month_start_day'] + i)
        for i in range(blank_days_count):  # 反
            #这里是一整行添加 ，倒过来的一定要减 2！！！！！
            for j in range (s_param['table_col_len']):
                col_0 = table.cell(-i - 2,j).paragraphs[0]
                col_0.text = '\\'
    elif d_h_param['genTable_switch'] == s_param['gentable_next_mode_num']:
        for i in range(s_param['days_len']):
            col_0 = table.cell(i + 1, 0).paragraphs[0]
            col_0.text = str(d_h_param['next_month']) + '.' + '%s' % (s_param['genTable_next_month_start_day'] + i)
    elif d_h_param['genTable_switch'] == s_param['gentable_repaire_early_mode_num']:
        for i in range(s_param['days_len']):
            col_0 = table.cell(i + 1, 0).paragraphs[0]
            col_0.text = str(d_h_param['month'] )+ '.' + '%s' % (s_param['genTable_next_month_start_day'] + i)
    else:
        ex = Exception('No other modes now!')
        raise ex

    document.save(s_param['file_path'] + '%s' % d_h_param['genTable_headline_str'] + s_param['file_tpye'])

def static_score():
    #再最后一项填入得分
    f = Document(s_param['file_path'] + '%s' % d_h_param['current_headline_str'] + s_param['file_tpye'])
    table = f.tables[0]
    score=0
    for i in range(s_param['days_len']):
        for j in range(s_param['things_len']):
            if table.cell(i+1,j+1).text == '✔':
                score+=1
    table.cell(s_param['days_len'],s_param['things_len']).paragraphs[0].text = score
    table.cell(s_param['days_len'],s_param['things_len']).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    f.save(s_param['file_path'] + '%s' % d_h_param['current_headline_str'] + s_param['file_tpye'])  # 需要保存！


if __name__ == '__main__':
    pass