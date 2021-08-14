#!/usr/bin/env python 
# -*- coding: utf-8 -*-
# @Time    : 2021/8/12 21:36
# @Author  : hanyou
# @Software: PyCharm
from docx import Document

from clock import *

#
# creat_window()


#读取
# f= Document('test.docx')
# for line in f.paragraphs:
#     print(line.text)

path = 'test.docx' #文件路径
document = Document(path) #读入文件

tables = document.tables #获取文件中的表格集
table = tables[0]#获取文件中的第一个表格
for i in range(0,13):#从表格第1行开始循环读取表格数据
    result = table.cell(0,i).text + table.cell(1,i).text+table.cell(2,i).text + table.cell(3,i).text
    #cell(i,0)表示第(i+1)行第1列数据，以此类推
    print(result)
